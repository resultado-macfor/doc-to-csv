import streamlit as st
import google.generativeai as genai
import pandas as pd
import os
from datetime import datetime
import tempfile
import docx
import io
import csv
import json
import re
from PIL import Image, ImageDraw, ImageFont
import time

# Configuração
st.set_page_config(page_title="Extrator de Cultivares", page_icon="🌱", layout="wide")
st.title("🌱 Extrator de Cultivares - DOCX para Google Sheets")

# API Key
gemini_api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GEM_API_KEY")
if not gemini_api_key:
    st.error("Configure GEMINI_API_KEY ou GEM_API_KEY")
    st.stop()

try:
    genai.configure(api_key=gemini_api_key)
    modelo_visao = genai.GenerativeModel("gemini-2.0-flash-exp")
    modelo_texto = genai.GenerativeModel("gemini-1.5-flash")
except Exception as e:
    st.error(f"Erro ao configurar Gemini: {str(e)}")
    st.stop()

# COLUNAS EXATAS conforme o template esperado
COLUNAS_EXATAS = [
    "Cultura", "Nome do produto", "NOME TÉCNICO/ REG", "Descritivo para SEO", 
    "Fertilidade", "Grupo de maturação", "Lançamento", "Slogan", "Tecnologia", 
    "Região (por extenso)", "Estado (por extenso)", "Ciclo", "Finalidade", 
    "URL da imagem do mapa", "Número do ícone", "Titulo icone 1", "Descrição Icone 1", 
    "Número do ícone", "Titulo icone 2", "Descrição Icone 2", "Número do ícone", 
    "Titulo icone 3", "Descrição Icone 3", "Número do ícone", "Título icone 4", 
    "Descrição Icone 4", "Número do ícone", "Título icone 5", "Descrição Icone 5", 
    "Exigência à fertilidade", "Grupo de maturidade", "PMS MÉDIO", "Tipo de crescimento", 
    "Cor da flor", "Cor da pubescência", "Cor do hilo", "Cancro da haste", 
    "Pústula bacteriana", "Nematoide das galhas - M. javanica", 
    "Nematóide de Cisto (Raça 3", "Nematóide de Cisto (Raça 9)", 
    "Nematóide de Cisto (Raça 10", "Nematóide de Cisto (Raça 14)", 
    "Fitóftora (Raça 1)", "Recomendações", "Resultado 1 - Nome", "Resultado 1 - Local", 
    "Resultado 1", "Resultado 2 - Nome", "Resultado 2 - Local", "Resultado 2", 
    "Resultado 3 - Nome", "Resultado 3 - Local", "Resultado 3", "Resultado 4 - Nome", 
    "Resultado 4 - Local", "Resultado 4", "Resultado 5 - Nome", "Resultado 5 - Lcal", 
    "Resultado 5", "Resultado 6 - Nome", "Resultado 6 - Local", "Resultado 6", 
    "Resultado 7 - Nome", "Resultado 7 - Local", "Resultado 7", "REC", "UF", 
    "Região", "Mês 1", "Mês 2", "Mês 3", "Mês 4", "Mês 5", "Mês 6", "Mês 7", 
    "Mês 8", "Mês 9", "Mês 10", "Mês 11", "Mês 12"
]

# Versão alternativa para mapeamento interno
COLUNAS_INTERNAS = [
    "Cultura", "Nome do produto", "NOME TÉCNICO/ REG", "Descritivo para SEO", 
    "Fertilidade", "Grupo de maturação", "Lançamento", "Slogan", "Tecnologia", 
    "Região (por extenso)", "Estado (por extenso)", "Ciclo", "Finalidade", 
    "URL da imagem do mapa", "Número do ícone", "Titulo icone 1", "Descrição Icone 1", 
    "Número do ícone", "Titulo icone 2", "Descrição Icone 2", "Número do ícone", 
    "Titulo icone 3", "Descrição Icone 3", "Número do ícone", "Título icone 4", 
    "Descrição Icone 4", "Número do ícone", "Título icone 5", "Descrição Icone 5", 
    "Exigência à fertilidade", "Grupo de maturidade", "PMS MÉDIO", "Tipo de crescimento", 
    "Cor da flor", "Cor da pubescência", "Cor do hilo", "Cancro da haste", 
    "Pústula bacteriana", "Nematoide das galhas - M. javanica", 
    "Nematóide de Cisto (Raça 3)", "Nematóide de Cisto (Raça 9)", 
    "Nematóide de Cisto (Raça 10)", "Nematóide de Cisto (Raça 14)", 
    "Fitóftora (Raça 1)", "Recomendações", "Resultado 1 - Nome", "Resultado 1 - Local", 
    "Resultado 1", "Resultado 2 - Nome", "Resultado 2 - Local", "Resultado 2", 
    "Resultado 3 - Nome", "Resultado 3 - Local", "Resultado 3", "Resultado 4 - Nome", 
    "Resultado 4 - Local", "Resultado 4", "Resultado 5 - Nome", "Resultado 5 - Lcal", 
    "Resultado 5", "Resultado 6 - Nome", "Resultado 6 - Local", "Resultado 6", 
    "Resultado 7 - Nome", "Resultado 7 - Local", "Resultado 7", "REC", "UF", 
    "Região", "Mês 1", "Mês 2", "Mês 3", "Mês 4", "Mês 5", "Mês 6", "Mês 7", 
    "Mês 8", "Mês 9", "Mês 10", "Mês 11", "Mês 12"
]

# Mapeamento entre nomes internos e nomes exatos do template
MAPEAMENTO_COLUNAS = {
    "Nematóide de Cisto (Raça 3)": "Nematóide de Cisto (Raça 3",
    "Nematóide de Cisto (Raça 10)": "Nematóide de Cisto (Raça 10",
    "Nematóide de Cisto (Raça 14)": "Nematóide de Cisto (Raça 14",
    "Pústula bacteriana": "Pústula bacteriana ",
    "Resultado 5 - Lcal": "Resultado 5 - Lcal"  # Note o typo intencional
}

# Inicializar session state
if 'df' not in st.session_state:
    st.session_state.df = pd.DataFrame(columns=COLUNAS_EXATAS)
if 'csv_content' not in st.session_state:
    st.session_state.csv_content = ""
if 'imagens' not in st.session_state:
    st.session_state.imagens = []
if 'texto' not in st.session_state:
    st.session_state.texto = ""

# Função 1: Converter DOCX para texto diretamente
def docx_para_texto(docx_bytes):
    """Converte DOCX diretamente para texto preservando estrutura"""
    try:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(docx_bytes)
            docx_path = tmp.name
        
        doc = docx.Document(docx_path)
        
        # Extrair todo o texto com estrutura
        textos = []
        
        # Extrair parágrafos
        for para in doc.paragraphs:
            if para.text.strip():
                textos.append(para.text.strip())
        
        # Extrair tabelas com formatação
        for table in doc.tables:
            tabela_texto = []
            for i, row in enumerate(table.rows):
                linha = []
                for cell in row.cells:
                    if cell.text.strip():
                        linha.append(cell.text.strip())
                if linha:
                    tabela_texto.append(" | ".join(linha))
            if tabela_texto:
                textos.append("--- TABELA ---")
                textos.extend(tabela_texto)
                textos.append("--- FIM TABELA ---")
        
        texto_completo = "\n".join(textos)
        os.unlink(docx_path)
        
        return texto_completo
        
    except Exception as e:
        st.error(f"Erro na conversão DOCX: {str(e)}")
        return ""

# Função 2: Melhorar transcrição com prompts específicos
def transcrever_texto(texto_original):
    """Melhora e organiza o texto extraído para análise"""
    
    prompt = f"""
    Você é um especialista em processamento de documentos técnicos de agricultura.
    
    TEXTO ORIGINAL EXTRAÍDO DO DOCUMENTO:
    {texto_original}
    
    SUA TAREFA:
    1. Reorganize este texto mantendo TODA a informação
    2. Identifique e separe claramente cada cultivar/produto
    3. Para cada cultivar, estruture as informações em seções:
       - Identificação (nome, cultura, tecnologia)
       - Características técnicas (ciclo, fertilidade, grupo de maturação)
       - Resistências a doenças
       - Resultados de produtividade
       - Recomendações
       - Regiões/Estados recomendados
    
    4. PRESERVE todos os dados numéricos, nomes, siglas, resultados
    5. Use marcadores claros como "=== CULTIVAR: [NOME] ===" para separar
    
    Retorne apenas o texto reorganizado e estruturado.
    """
    
    try:
        response = modelo_texto.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Erro na estruturação do texto: {str(e)}")
        return texto_original

# Função 3: Extrair dados com prompt mais específico
def extrair_dados_para_csv(texto_estruturado):
    """Extrai dados do texto para o formato CSV com prompt detalhado"""
    
    prompt = f"""
    ANALISE O TEXTO ESTRUTURADO ABAIXO E EXTRAIA TODOS OS DADOS PARA PREENCHER A PLANILHA DE CULTIVARES.

    TEXTO ESTRUTURADO:
    {texto_estruturado}

    **INSTRUÇÕES CRÍTICAS:**

    1. **IDENTIFIQUE CADA CULTIVAR ÚNICA** no documento. Cada cultivar deve ser uma linha na planilha.

    2. **USE O EXEMPLO ABAIXO COMO REFERÊNCIA** para entender como preencher as colunas:

    EXEMPLO DE LINHA PREENCHIDA (EXATAMENTE como deve aparecer no CSV):
    Cultura: Soja
    Nome do produto: NS7524IPRO
    NOME TÉCNICO/ REG: 
    Descritivo para SEO: 
    Fertilidade: Alto
    Grupo de maturação: 7.5
    Lançamento: Sim
    Slogan: Excelente performance produtiva com múltipla resistência a nematoides de cisto
    Tecnologia: IPRO
    Região (por extenso): Sul, Sudeste
    Estado (por extenso): Santa Catarina, Paraná
    Ciclo: Precoce
    Finalidade: Grãos
    URL da imagem do mapa: Ex: https://www.niderasementes.com.br/wp-content/uploads/2025/12/mapa_soja_niderasementes-1000x1000.jpg
    Número do ícone: 1
    Titulo icone 1: Ex: Alto retorno ao investimento
    Descrição Icone 1: Altíssimo potencial produtivo; Indicada para alta tecnologia: no melhor talhão e com melhor manejo
    Número do ícone: 
    Titulo icone 2: Ex: Facilidade do plantio a colheita
    Descrição Icone 2: Excelente estabelecimento inicial de plantas; Arquitetura de planta que facilita o manejo; Bom comportamento em relação ao acamamento.
    Número do ícone: 
    Titulo icone 3: 
    Descrição Icone 3: 
    Número do ícone: 
    Título icone 4: 
    Descrição Icone 4: 
    Número do ícone: 
    Título icone 5: 
    Descrição Icone 5: 
    Exigência à fertilidade: Médio e alto
    Grupo de maturidade: 7.7 M3 | 7.8 M4 | 7.8 M5
    PMS MÉDIO: 150G
    Tipo de crescimento: Semideterminado
    Cor da flor: Roxa
    Cor da pubescência: Marrom média
    Cor do hilo: Preto
    Cancro da haste: R
    Pústula bacteriana: MR
    Nematoide das galhas - M. javanica: R
    Nematóide de Cisto (Raça 3: R
    Nematóide de Cisto (Raça 9): MR
    Nematóide de Cisto (Raça 10: MR
    Nematóide de Cisto (Raça 14): MR
    Fitóftora (Raça 1): MR
    Recomendações: Pode haver variação no ciclo (dias) devido às condições edafoclimáticas...
    Resultado 1 - Nome: Fazenda Planalto
    Resultado 1 - Local: Costa Rica - MS
    Resultado 1: 106,0 sc/ha
    Resultado 2 - Nome: Clodemir Paholski
    Resultado 2 - Local: Cristalina - GO
    Resultado 2: 85,0 sc/ha
    Resultado 3 - Nome: Centro Sul Consultoria
    Resultado 3 - Local: Formosa – GO
    Resultado 3: 84,5 sc/ha
    Resultado 4 - Nome: Antério Mânica
    Resultado 4 - Local: Unaí - MG
    Resultado 4: 84,0 sc/ha
    Resultado 5 - Nome: Cislei Ribeiro dos Santos
    Resultado 5 - Local: Bonfinópolis de Minas - MG
    Resultado 5: 84,0 sc/ha
    Resultado 6 - Nome: Djonas Kogler
    Resultado 6 - Local: Formoso - MG
    Resultado 6: 81,0 sc/ha
    Resultado 7 - Nome: Cerrado Consultoria
    Resultado 7 - Local: Unaí - MG
    Resultado 7: 79,0 sc/ha
    REC: 202
    UF: RS, SC, PR, SP
    Região: Sul, Sudeste
    Mês 1: NR
    Mês 2: NR
    Mês 3: 180-260
    Mês 4: 180-260
    Mês 5: 180-260
    Mês 6: 180-260
    Mês 7: 180-260
    Mês 8: 180-260
    Mês 9: 180-260
    Mês 10: 180-260
    Mês 11: 180-260
    Mês 12: NR

    3. **PARA CADA CULTIVAR IDENTIFICADA**, preencha TODAS as 81 colunas listadas abaixo.
    **USE OS NOMES DE COLUNAS EXATAMENTE COMO ESTÃO LISTADOS AQUI:**

    LISTA EXATA DAS 81 COLUNAS (COPIE E COLE ESTES NOMES):
    Cultura, Nome do produto, NOME TÉCNICO/ REG, Descritivo para SEO, Fertilidade, Grupo de maturação, Lançamento, Slogan, Tecnologia, Região (por extenso), Estado (por extenso), Ciclo, Finalidade, URL da imagem do mapa, Número do ícone, Titulo icone 1, Descrição Icone 1, Número do ícone, Titulo icone 2, Descrição Icone 2, Número do ícone, Titulo icone 3, Descrição Icone 3, Número do ícone, Título icone 4, Descrição Icone 4, Número do ícone, Título icone 5, Descrição Icone 5, Exigência à fertilidade, Grupo de maturidade, PMS MÉDIO, Tipo de crescimento, Cor da flor, Cor da pubescência, Cor do hilo, Cancro da haste, Pústula bacteriana , Nematoide das galhas - M. javanica, Nematóide de Cisto (Raça 3, Nematóide de Cisto (Raça 9), Nematóide de Cisto (Raça 10, Nematóide de Cisto (Raça 14), Fitóftora (Raça 1), Recomendações, Resultado 1 - Nome, Resultado 1 - Local, Resultado 1, Resultado 2 - Nome, Resultado 2 - Local, Resultado 2, Resultado 3 - Nome, Resultado 3 - Local, Resultado 3, Resultado 4 - Nome, Resultado 4 - Local, Resultado 4, Resultado 5 - Nome, Resultado 5 - Lcal, Resultado 5, Resultado 6 - Nome, Resultado 6 - Local, Resultado 6, Resultado 7 - Nome, Resultado 7 - Local, Resultado 7, REC, UF, Região, Mês 1, Mês 2, Mês 3, Mês 4, Mês 5, Mês 6, Mês 7, Mês 8, Mês 9, Mês 10, Mês 11, Mês 12

    4. **FORMATO DE RESPOSTA:**
    Retorne APENAS um array JSON onde cada objeto representa uma cultivar.
    Cada objeto deve ter EXATAMENTE 81 propriedades com os nomes das colunas acima.
    USE OS NOMES DE COLUNAS EXATAMENTE COMO ESTÃO NA LISTA ACIMA.

    5. **REGRA DE PREENCHIMENTO:**
    - Para informações não encontradas, use "NR"
    - Para múltiplos valores (como estados), separe com vírgula: "Santa Catarina, Paraná"
    - Para resultados de produtividade, mantenha o formato: "106,0 sc/ha"
    - Para resistências, use siglas: R (Resistente), MR (Moderadamente Resistente), S (Suscetível)
    - Para meses de plantio, use formato "180-260" ou "NR"
    - Note que algumas colunas tem nomes incompletos como "Nematóide de Cisto (Raça 3" (sem fechar parênteses)

    6. **ATENÇÃO AOS DETALHES:**
    - "Número do ícone" se repete 5 vezes (para cada ícone)
    - "Resultado 5 - Lcal" tem typo (é "Lcal" não "Local")
    - "Pústula bacteriana " tem espaço no final
    - Algumas colunas de nematóide não fecham parênteses

    AGORA ANALISE O TEXTO E EXTRAIA OS DADOS PARA O ARRAY JSON COM AS 81 COLUNAS EXATAS:
    """
    
    try:
        with st.spinner("🔍 Extraindo dados de cada cultivar..."):
            response = modelo_texto.generate_content(prompt)
            resposta = response.text.strip()
            
            # Limpar resposta
            resposta_limpa = resposta.replace('```json', '').replace('```', '').replace('JSON', '').strip()
            
            # Tentar encontrar JSON
            try:
                # Tentar parse direto
                dados = json.loads(resposta_limpa)
                
                # Verificar se é lista
                if isinstance(dados, list):
                    return dados
                elif isinstance(dados, dict):
                    # Se for objeto único, colocar em lista
                    return [dados]
                else:
                    st.warning("Formato de resposta inesperado")
                    return []
                    
            except json.JSONDecodeError:
                # Tentar extrair JSON com regex
                json_match = re.search(r'(\[.*\])', resposta_limpa, re.DOTALL)
                if json_match:
                    json_str = json_match.group(1)
                    dados = json.loads(json_str)
                    return dados
                
                # Tentar objeto único
                obj_match = re.search(r'(\{.*\})', resposta_limpa, re.DOTALL)
                if obj_match:
                    json_str = obj_match.group(1)
                    dados = [json.loads(json_str)]
                    return dados
                
                st.warning("Não foi possível extrair JSON da resposta")
                st.info(f"Resposta recebida: {resposta_limpa[:500]}...")
                return []
            
    except Exception as e:
        st.error(f"Erro na extração de dados: {str(e)}")
        st.info(f"Resposta que causou erro: {resposta[:1000]}")
        return []

# Função 4: Criar DataFrame com validação e mapeamento correto
def criar_dataframe(dados):
    """Cria DataFrame a partir dos dados extraídos com validação"""
    if not dados or not isinstance(dados, list):
        return pd.DataFrame(columns=COLUNAS_EXATAS)
    
    linhas = []
    for idx, item in enumerate(dados):
        if isinstance(item, dict):
            linha = {}
            
            # Primeiro, coletar todos os valores com nomes flexíveis
            valores_coletados = {}
            
            for chave, valor in item.items():
                chave_limpa = chave.strip()
                if isinstance(valor, str):
                    valor_limpo = valor.strip()
                else:
                    valor_limpo = str(valor).strip()
                
                # Mapear para nomes internos primeiro
                for coluna_interna in COLUNAS_INTERNAS:
                    if chave_limpa == coluna_interna or chave_limpa in coluna_interna:
                        valores_coletados[coluna_interna] = valor_limpo
                        break
            
            # Agora preencher as colunas exatas
            for coluna_exata in COLUNAS_EXATAS:
                valor_final = "NR"
                
                # Tentar encontrar valor correspondente
                for coluna_interna, valor in valores_coletados.items():
                    # Mapeamento específico para colunas problemáticas
                    if coluna_exata in MAPEAMENTO_COLUNAS and MAPEAMENTO_COLUNAS.get(coluna_interna) == coluna_exata:
                        valor_final = valor
                        break
                    # Verificar correspondência geral
                    elif coluna_exata == coluna_interna or coluna_exata in coluna_interna:
                        valor_final = valor
                        break
                
                linha[coluna_exata] = valor_final
            
            # Validar linha mínima
            if linha.get("Nome do produto", "NR") != "NR" and linha.get("Cultura", "NR") != "NR":
                linhas.append(linha)
            else:
                st.warning(f"Linha {idx+1} ignorada: falta nome do produto ou cultura")
    
    if linhas:
        df = pd.DataFrame(linhas, columns=COLUNAS_EXATAS)
        return df
    else:
        return pd.DataFrame(columns=COLUNAS_EXATAS)

# Função 5: Gerar CSV formatado corretamente com colunas exatas
def gerar_csv_para_gsheets(df):
    """Gera CSV formatado para Google Sheets com tratamento especial"""
    if df.empty:
        return ""
    
    output = io.StringIO()
    
    # Configurar writer para preservar formato brasileiro
    writer = csv.writer(output, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    
    # Escrever cabeçalho EXATO
    writer.writerow(COLUNAS_EXATAS)
    
    # Escrever dados
    for _, row in df.iterrows():
        linha = []
        for col in COLUNAS_EXATAS:
            valor = str(row.get(col, "NR")).strip()
            
            # Tratamentos especiais
            if valor == "":
                valor = "NR"
            elif valor in ["nan", "None", "null", "NaN"]:
                valor = "NR"
            
            # Preservar vírgulas decimais (formato brasileiro)
            if "sc/ha" in valor or "," in valor and valor.replace(",", "").replace(".", "").isdigit():
                # Manter vírgula decimal
                valor = valor
            
            linha.append(valor)
        
        writer.writerow(linha)
    
    csv_content = output.getvalue()
    
    return csv_content

# Interface principal
def main():
    st.sidebar.header("📤 Upload do Documento")
    
    uploaded_file = st.sidebar.file_uploader(
        "Carregue um arquivo DOCX técnico de cultivares:",
        type=["docx"],
        help="Documento contendo informações sobre cultivares/soja com tabelas, características e resultados"
    )
    
    if uploaded_file:
        st.sidebar.info(f"**Arquivo:** {uploaded_file.name}")
        st.sidebar.info(f"**Tamanho:** {uploaded_file.size/1024:.1f} KB")
        
        # Mostrar as colunas esperadas
        with st.sidebar.expander("📋 Ver 81 colunas esperadas", expanded=False):
            st.write("Colunas exatas que serão geradas:")
            for i, col in enumerate(COLUNAS_EXATAS[:20], 1):
                st.write(f"{i}. {col}")
            st.write(f"... e mais {len(COLUNAS_EXATAS)-20} colunas")
        
        col1, col2 = st.sidebar.columns(2)
        
        with col1:
            if st.button("🚀 Processar Documento", type="primary", use_container_width=True):
                # Limpar estado anterior
                st.session_state.imagens = []
                st.session_state.texto = ""
                st.session_state.df = pd.DataFrame(columns=COLUNAS_EXATAS)
                st.session_state.csv_content = ""
                
                with st.spinner("📄 Convertendo DOCX para texto..."):
                    # PASSO 1: Converter DOCX para texto
                    texto_original = docx_para_texto(uploaded_file.getvalue())
                    
                    if not texto_original or len(texto_original) < 100:
                        st.error("Documento muito curto ou vazio")
                        return
                    
                    st.success(f"✅ Texto extraído ({len(texto_original):,} caracteres)")
                
                with st.spinner("🧹 Estruturando e organizando texto..."):
                    # PASSO 2: Melhorar estrutura do texto
                    texto_estruturado = transcrever_texto(texto_original)
                    st.session_state.texto = texto_estruturado
                    st.success("✅ Texto estruturado")
                
                with st.spinner("📊 Extraindo dados para 81 colunas exatas..."):
                    # PASSO 3: Extrair dados estruturados
                    dados = extrair_dados_para_csv(texto_estruturado)
                    
                    if dados:
                        st.info(f"✅ {len(dados)} cultivar(s) identificada(s)")
                        
                        # PASSO 4: Criar DataFrame
                        df = criar_dataframe(dados)
                        st.session_state.df = df
                        
                        if not df.empty:
                            st.success(f"✅ DataFrame criado com {len(df)} linha(s) e {len(df.columns)} coluna(s)")
                            
                            # Verificar colunas geradas
                            colunas_geradas = list(df.columns)
                            colunas_corretas = all(col in COLUNAS_EXATAS for col in colunas_geradas)
                            
                            if colunas_corretas and len(colunas_geradas) == 81:
                                st.success("✅ Todas as 81 colunas exatas foram geradas!")
                            else:
                                st.warning(f"⚠️ {len(colunas_geradas)} colunas geradas (esperado: 81)")
                            
                            # PASSO 5: Gerar CSV
                            csv_content = gerar_csv_para_gsheets(df)
                            st.session_state.csv_content = csv_content
                            st.success("✅ CSV gerado para Google Sheets")
                        else:
                            st.warning("⚠️ DataFrame vazio após processamento")
                    else:
                        st.error("❌ Nenhum dado extraído do documento")
        
        with col2:
            if st.button("🔄 Limpar Tudo", use_container_width=True):
                st.session_state.imagens = []
                st.session_state.texto = ""
                st.session_state.df = pd.DataFrame(columns=COLUNAS_EXATAS)
                st.session_state.csv_content = ""
                st.rerun()
        
        # Mostrar resultados
        df = st.session_state.df
        
        # Verificar se temos dados para mostrar
        if df is not None and not df.empty:
            st.header("📊 Resultados - Pronto para Google Sheets")
            
            # Verificação das colunas
            st.subheader("✅ Verificação das Colunas Geradas")
            
            col_ver1, col_ver2, col_ver3 = st.columns(3)
            
            with col_ver1:
                st.metric("Colunas Geradas", len(df.columns))
            
            with col_ver2:
                st.metric("Colunas Esperadas", 81)
            
            with col_ver3:
                if len(df.columns) == 81:
                    st.success("✅ OK")
                else:
                    st.error(f"❌ Faltam {81 - len(df.columns)}")
            
            # Mostrar comparação de colunas
            with st.expander("🔍 Comparar colunas geradas vs esperadas", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**Colunas Geradas:**")
                    for i, col in enumerate(df.columns[:40], 1):
                        st.write(f"{i}. {col}")
                
                with col2:
                    st.write("**Colunas Esperadas (primeiras 40):**")
                    for i, col in enumerate(COLUNAS_EXATAS[:40], 1):
                        st.write(f"{i}. {col}")
            
            # Estatísticas
            st.subheader("📈 Estatísticas dos Dados")
            
            col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
            with col_stat1:
                st.metric("Cultivares", len(df))
            with col_stat2:
                if 'Cultura' in df.columns:
                    culturas_unicas = df['Cultura'].nunique()
                    st.metric("Culturas", culturas_unicas)
            with col_stat3:
                tecnologias_unicas = df['Tecnologia'].nunique() if 'Tecnologia' in df.columns else 0
                st.metric("Tecnologias", tecnologias_unicas)
            with col_stat4:
                colunas_preenchidas = sum([1 for col in df.columns if df[col].astype(str).str.contains('NR').mean() < 0.8])
                st.metric("Colunas Preenchidas", colunas_preenchidas)
            
            # Visualização dos dados
            st.subheader("👁️ Visualização dos Dados")
            
            # Selecionar colunas para mostrar
            colunas_importantes = [
                'Cultura', 'Nome do produto', 'Tecnologia', 
                'Grupo de maturação', 'Ciclo', 'Fertilidade',
                'Estado (por extenso)', 'PMS MÉDIO'
            ]
            
            colunas_disponiveis = [c for c in colunas_importantes if c in df.columns]
            
            if colunas_disponiveis:
                st.dataframe(df[colunas_disponiveis], use_container_width=True, height=300)
            
            # Botão para visualizar todas as colunas
            with st.expander("👁️ Ver todas as 81 colunas", expanded=False):
                st.dataframe(df, use_container_width=True, height=400)
            
            # Download
            st.subheader("📥 Download dos Arquivos")
            
            nome_base = uploaded_file.name.split('.')[0]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            col_dl1, col_dl2, col_dl3 = st.columns(3)
            
            with col_dl1:
                # CSV para Google Sheets
                if st.session_state.csv_content:
                    st.download_button(
                        label="📄 Baixar CSV (81 colunas exatas)",
                        data=st.session_state.csv_content.encode('utf-8'),
                        file_name=f"cultivares_{nome_base}_{timestamp}.csv",
                        mime="text/csv",
                        help="CSV com 81 colunas EXATAS conforme template",
                        use_container_width=True
                    )
            
            with col_dl2:
                # Excel
                if not df.empty:
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        df.to_excel(writer, index=False, sheet_name='Cultivares')
                    excel_data = excel_buffer.getvalue()
                    
                    st.download_button(
                        label="📊 Baixar Excel (.xlsx)",
                        data=excel_data,
                        file_name=f"cultivares_{nome_base}_{timestamp}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        help="Arquivo Excel para edição",
                        use_container_width=True
                    )
            
            with col_dl3:
                # JSON para depuração
                if st.session_state.texto:
                    st.download_button(
                        label="📝 Baixar Texto Estruturado",
                        data=st.session_state.texto,
                        file_name=f"texto_estruturado_{nome_base}_{timestamp}.txt",
                        mime="text/plain",
                        help="Texto estruturado para análise",
                        use_container_width=True
                    )
            
            # Preview do CSV
            with st.expander("🔍 Preview do CSV gerado", expanded=False):
                if st.session_state.csv_content:
                    # Mostrar cabeçalho
                    linhas = st.session_state.csv_content.split('\n')[:2]
                    st.code("\n".join(linhas), language="csv")
                    
                    # Mostrar contagem de colunas
                    if len(linhas) > 0:
                        cabecalho = linhas[0]
                        colunas_count = cabecalho.count(',') + 1
                        st.info(f"**{colunas_count} colunas no CSV**")
        
        elif df is not None and df.empty:
            st.info("📭 Nenhuma cultivar identificada no documento.")
        
        # Mostrar pipeline completo
        with st.expander("⚙️ Pipeline Completo", expanded=True):
            st.markdown(f"""
            ### 🔄 **Fluxo de Processamento:**
            
            1. **📤 DOCX Original**  
               → Extração direta de texto e tabelas
            
            2. **🧹 Estruturação com IA**  
               → Identificação de cultivares individuais  
               → Organização em seções lógicas
            
            3. **📊 Extração para 81 Colunas EXATAS**  
               → {len(COLUNAS_EXATAS)} colunas conforme template  
               → Mapeamento detalhado para cada campo  
               → Validação de dados  
            
            4. **📄 CSV Final**  
               → Colunas exatas conforme template  
               → Pronto para Google Sheets  
               → Cada linha = uma página no site
            
            **Status atual:**
            """)
            
            status_col1, status_col2, status_col3 = st.columns(3)
            
            with status_col1:
                if st.session_state.texto:
                    st.success("✅ Texto estruturado")
                else:
                    st.info("📭 Aguardando")
            
            with status_col2:
                if not st.session_state.df.empty:
                    st.success(f"✅ {len(st.session_state.df)} cultivar(s)")
                else:
                    st.info("📭 Aguardando")
            
            with status_col3:
                if st.session_state.csv_content:
                    st.success("✅ CSV pronto")
                else:
                    st.info("📭 Aguardando")
    
    else:
        # Tela inicial
        st.markdown("""
        ## 🌱 Extrator de Cultivares - Pipeline Completo
        
        ### 🎯 **Objetivo:**
        Transformar documentos técnicos de cultivares em planilhas estruturadas com **81 colunas EXATAS** para o site.
        
        ### 📋 **Template de Saída (81 Colunas EXATAS):**
        
        As colunas serão geradas **EXATAMENTE** conforme este template:
        
        1. **Cultura** 2. **Nome do produto** 3. **NOME TÉCNICO/ REG**  
        4. **Descritivo para SEO** 5. **Fertilidade** 6. **Grupo de maturação**  
        7. **Lançamento** 8. **Slogan** 9. **Tecnologia**  
        10. **Região (por extenso)** 11. **Estado (por extenso)** 12. **Ciclo**  
        ... e mais 69 colunas específicas.
        
        ### ✅ **Garantia:**
        - **81 colunas EXATAS** conforme template
        - **Mesmos nomes, mesma ordem**
        - **CSV pronto para importação** no Google Sheets
        - **Cada linha** = uma cultivar para o site
        
        ### 📤 **Como usar:**
        1. Carregue um documento DOCX na barra lateral
        2. Clique em **"Processar Documento"**
        3. Baixe o **CSV com 81 colunas exatas**
        4. Importe no **Google Sheets** - funcionará perfeitamente!
        
        **Pronto para começar? Carregue seu primeiro documento!**
        """)

if __name__ == "__main__":
    main()
